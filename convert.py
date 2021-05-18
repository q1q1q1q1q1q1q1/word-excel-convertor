import os
import docx
import xlsxwriter


class CUSTOM_TYPES:
    class ALIGNMENT:
        LEFT = 0
        CENTER = 1
        JUSTIFY = 3
        RIGHT = 2


class ParseWord:
    """analyzing the word.doc/docx syntax"""
    def __init__(self, path: str):
        assert len(path) > 0, "File path not specified"
        self._parsedText: [{"str"}] = []
        self._path = path

    def parse(self):
        doc = docx.Document(self._path)
        for paragraph in doc.paragraphs:
            self._parsedText.append([
                {
                    "text": paragraph.text,
                    "properties": {
                        "alignment": paragraph.alignment,
                        "indents": self.__getIndents(paragraph)
                    }

                }
            ])
        return self._parsedText

    def __getIndents(self, paragraph) -> {str}:
        """returning indents in pt format"""
        formatting = paragraph.paragraph_format
        before = formatting.space_before
        after = formatting.space_after
        left = formatting.left_indent
        right = formatting.right_indent
        first_line = formatting.first_line_indent

        if before is not None:
            before = before.pt

        if after is not None:
            after = after.pt

        if left is not None:
            left = left.pt

        if right is not None:
            right = right.pt

        if first_line is not None:
            first_line = first_line.pt

        return {
            "before": before,
            "after": after,
            "left": left,
            "right": right,
            "first_line": first_line
        }


class ImportInExcel:
    """filling in an excel file"""
    def __init__(self, text_for_import: [{str}]):
        assert len(text_for_import) > 0, "Error reading text from a word file"
        self.__text_for_import = text_for_import

        #maximum number of characters per cell
        self.__cell_capacity = 8

    def insert(self):
        workbook = xlsxwriter.Workbook('output.xlsx')
        worksheet = workbook.add_worksheet()
        worksheet.write(1, 1, "", )
        curr_line = 0
        for paragraph in self.__text_for_import:
            column = self.__findOutAlignment(alignment = paragraph[0]["properties"]["alignment"])
            worksheet.write(curr_line, column, paragraph[0]["text"])
            curr_line += 1

        workbook.close()

    def __findOutAlignment(self, alignment: CUSTOM_TYPES.ALIGNMENT):
        column = 0
        if alignment is not None:
            if alignment == CUSTOM_TYPES.ALIGNMENT.CENTER:
                column = 3
            elif alignment == CUSTOM_TYPES.ALIGNMENT.JUSTIFY:
                column = 0
            elif alignment == CUSTOM_TYPES.ALIGNMENT.LEFT:
                column = 0
            elif alignment == CUSTOM_TYPES.ALIGNMENT.RIGHT:
                column = 8
        return column


def main():
    # doc = docx.Document(os.path.join(os.path.abspath(os.path.dirname(__file__)), 'text.docx'))
    pars = ParseWord(os.path.join(os.path.abspath(os.path.dirname(__file__)), 'cwt4.docx'))
    parsedText = pars.parse()
    importer = ImportInExcel(parsedText)
    importer.insert()


if __name__ == '__main__':
    main()
